"""
generate_dfd_word.py
Generates SMART_DFD.docx — the full DFD documentation for the SMART capstone system.
Run: python docs/generate_dfd_word.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(table):
    """Add thin borders to every cell in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{edge}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "BFBFBF")
                tcBorders.append(border)
            tcPr.append(tcBorders)


BLUE       = "1F3864"   # dark navy – headings
LIGHT_BLUE = "D6E4F0"  # cell header bg
ACCENT     = "2E75B6"  # table header row
WHITE      = "FFFFFF"
MONO       = "Courier New"
BODY       = "Calibri"


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = BODY


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    run.font.name = BODY


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = BODY
    # bottom border line
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run.font.name = BODY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = BODY
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = BODY
    p.paragraph_format.left_indent = Cm(1 + level * 0.75)


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = MONO
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, ACCENT)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
                run.font.name = BODY

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F2F7FB" if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = BODY

    # column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    set_cell_border(table)
    doc.add_paragraph()   # spacing after table


# ── document ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────────────────
    add_title(doc, "SMART System")
    add_title(doc, "Data Flow Diagram (DFD)")
    doc.add_paragraph()
    add_subtitle(doc, "School Management and Academic Records Technology")
    add_subtitle(doc, "Capstone Project — System Documentation")
    add_subtitle(doc, "Date: May 2026")
    doc.add_page_break()

    # ── Section 1: What is a DFD ────────────────────────────────────────────
    add_h1(doc, "1. What is a Data Flow Diagram (DFD)?")
    add_body(doc,
        "A Data Flow Diagram (DFD) is a graphical representation that shows how data "
        "moves through a system. It is widely used in system analysis and documentation "
        "to communicate how a system works to both technical and non-technical audiences.")
    add_body(doc, "A DFD answers four key questions:")
    add_bullet(doc, "Where does data come from? → External Entities (sources)")
    add_bullet(doc, "What happens to data? → Processes")
    add_bullet(doc, "Where is data stored? → Data Stores")
    add_bullet(doc, "Where does data go? → External Entities (destinations)")

    add_h2(doc, "1.1 DFD Levels")
    add_table(doc,
        ["Level", "Name", "Purpose"],
        [
            ["Level 0", "Context Diagram", "Shows the entire system as ONE process. All external entities and their connections to the system are visible."],
            ["Level 1", "Main DFD", "Breaks the system into its major processes (6–10 processes). Each process is numbered."],
            ["Level 2", "Sub-process DFD", "Expands each Level 1 process into smaller, more detailed steps."],
        ],
        col_widths=[1.0, 1.5, 4.0]
    )

    add_h2(doc, "1.2 DFD Symbols")
    add_table(doc,
        ["Symbol", "Name", "Description"],
        [
            ["Rectangle / Box", "External Entity", "A person, organization, or outside system that interacts with SMART."],
            ["Rounded Rectangle / Circle", "Process", "An action SMART performs on data. Labeled with a number and verb."],
            ["Open Rectangle (parallel lines)", "Data Store", "A database table or file where data is stored. Labeled D1, D2 …"],
            ["Arrow with label", "Data Flow", "Shows the movement of data between entities, processes, and stores."],
        ],
        col_widths=[1.8, 1.5, 3.5]
    )

    doc.add_page_break()

    # ── Section 2: System Overview ──────────────────────────────────────────
    add_h1(doc, "2. System Overview")
    add_body(doc,
        "SMART is a web-based school management system for a Junior High School "
        "(Grades 7–10). It is the central platform for managing academic records, "
        "class schedules, grading, attendance, and DepEd-mandated school form generation.")

    add_h2(doc, "2.1 Core Capabilities")
    for cap in [
        "User management with role-based access (Admin, Teacher, Registrar)",
        "Synchronization of teacher assignments from ATLAS (external system)",
        "Synchronization of student and enrollment data from EnrollPro (external system)",
        "Grade recording and quarterly grade computation using the DepEd formula",
        "Daily attendance tracking per class section",
        "Template-based generation of DepEd School Forms (SF1 through SF10)",
        "Electronic Class Record (ECR) template management and export",
        "System-wide audit logging for compliance and accountability",
    ]:
        add_bullet(doc, cap)

    add_h2(doc, "2.2 Key System Constraint")
    add_body(doc,
        "SMART connects to two external systems in READ-ONLY mode. It never writes, "
        "modifies, or deletes data on ATLAS or EnrollPro. All writes happen exclusively "
        "to SMART's own local database (smart_db, PostgreSQL).")

    add_table(doc,
        ["External System", "What SMART Reads From It", "SMART Writes Back?"],
        [
            ["ATLAS", "Teaching load assignments: which teacher handles which subject in which section and school year.", "NO — read-only"],
            ["EnrollPro", "Student records (LRN, name, gender), class sections, enrollments, and advisory assignments.", "NO — read-only"],
        ],
        col_widths=[1.5, 4.2, 1.0]
    )

    add_h2(doc, "2.3 User Roles")
    add_table(doc,
        ["Role", "Responsibilities"],
        [
            ["Admin", "Manages user accounts, configures school settings (school year, grading weights), monitors audit logs, triggers external data syncs."],
            ["Teacher", "Records student grades and attendance for their assigned class sections, exports Electronic Class Records (ECR)."],
            ["Registrar", "Uploads DepEd School Form templates, generates and exports all school forms (SF1–SF10) for any section."],
        ],
        col_widths=[1.5, 5.2]
    )

    doc.add_page_break()

    # ── Section 3: Level 0 ──────────────────────────────────────────────────
    add_h1(doc, "3. Level 0 — Context Diagram")
    add_body(doc,
        "The Context Diagram represents the entire SMART system as a single process. "
        "It shows all five external entities and the data that flows between them and the system.")

    add_code(doc, "                         [ATLAS]")
    add_code(doc, "                    (External System)")
    add_code(doc, "                          |")
    add_code(doc, "             Teaching Load Data (auto-sync / 30 min)")
    add_code(doc, "                          |")
    add_code(doc, "                          v")
    add_code(doc, "  [ADMIN]  --commands/settings-->  +---------------------+")
    add_code(doc, "  [ADMIN]  <--reports/audit logs-- |                     |")
    add_code(doc, "                                   |                     |")
    add_code(doc, "  [TEACHER] --grades/attendance-->  |      S M A R T      |")
    add_code(doc, "  [TEACHER] <--class records------  |                     |")
    add_code(doc, "                                   |                     |")
    add_code(doc, "  [REGISTRAR] --templates/req-->   |                     |")
    add_code(doc, "  [REGISTRAR] <--school forms----   +---------------------+")
    add_code(doc, "                                             |")
    add_code(doc, "                          Student/Enrollment Data (auto-sync / 30 min)")
    add_code(doc, "                                             |")
    add_code(doc, "                                             v")
    add_code(doc, "                                    [EnrollPro]")
    add_code(doc, "                            (External Enrollment System)")
    doc.add_paragraph()

    add_h2(doc, "3.1 External Entities Summary")
    add_table(doc,
        ["Entity", "Type", "Sends TO SMART", "Receives FROM SMART"],
        [
            ["ATLAS", "External System", "Teaching load assignments (teacher → subject → section)", "Nothing (read-only)"],
            ["EnrollPro", "External System", "Students, Sections, Enrollments, Adviser mappings", "Nothing (read-only)"],
            ["Admin", "Human User", "Login credentials, user commands, school settings, sync triggers", "User list, audit logs, sync status, reports"],
            ["Teacher", "Human User", "Login credentials, grade entries, attendance records", "Class lists, computed grades, class record files"],
            ["Registrar", "Human User", "Login credentials, SF template files, report requests", "Generated School Forms (SF1–SF10), student reports"],
        ],
        col_widths=[1.2, 1.2, 2.6, 2.0]
    )

    doc.add_page_break()

    # ── Section 4: Level 1 ──────────────────────────────────────────────────
    add_h1(doc, "4. Level 1 — Main Processes")
    add_body(doc,
        "Level 1 breaks SMART into its seven main functional processes. Each process "
        "is numbered and shows the data stores it reads from or writes to.")

    processes = [
        ("1.0", "Authentication",
         "Validates user credentials, generates JWT session tokens, and logs login activity.",
         [("Input", "Username, password, role (from Admin/Teacher/Registrar)"),
          ("Reads", "D1: Users"),
          ("Writes", "D10: AuditLog"),
          ("Output", "JWT Bearer Token (session)")]
        ),
        ("2.0", "Data Synchronization",
         "Pulls and merges teaching load data from ATLAS and enrollment data from EnrollPro into SMART's database on a scheduled or manual basis.",
         [("Input", "ATLAS faculty assignments, EnrollPro students/sections/enrollments"),
          ("Reads", "(ATLAS API), (EnrollPro API)"),
          ("Writes", "D2: Teachers, D3: Students, D4: Sections, D5: Subjects, D6: ClassAssignments, D7: Enrollments"),
          ("Output", "Sync status, error logs")]
        ),
        ("3.0", "Grade Management",
         "Records student component scores, computes quarterly grades using the DepEd formula, and generates Electronic Class Record (ECR) files.",
         [("Input", "Written Work scores, Performance Task scores, Quarterly Assessment score"),
          ("Reads", "D3: Students, D6: ClassAssignments, D12: GradingConfig, D14: ECRTemplates"),
          ("Writes", "D8: Grades"),
          ("Output", "Computed quarterly grades, downloadable ECR Excel file")]
        ),
        ("4.0", "Attendance Management",
         "Records daily attendance per student per section and generates attendance summaries.",
         [("Input", "Student ID, date, attendance status (Present / Absent / Late / Excused)"),
          ("Reads", "D3: Students, D4: Sections"),
          ("Writes", "D9: Attendance"),
          ("Output", "Attendance summaries, SF2/SF4 report data")]
        ),
        ("5.0", "Template Management",
         "Handles upload and storage of DepEd School Form Excel templates and ECR templates.",
         [("Input", "Excel (.xlsx) template files uploaded by Registrar or Admin"),
          ("Writes", "D13: ExcelTemplates, D14: ECRTemplates"),
          ("Output", "Template list, confirmed upload")]
        ),
        ("6.0", "Report Generation",
         "Loads an Excel template, injects real data (students, grades, attendance, school info) into placeholder fields, and exports the completed form.",
         [("Input", "Form type request + section + school year parameters"),
          ("Reads", "D3: Students, D4: Sections, D8: Grades, D9: Attendance, D11: SystemSettings, D13: ExcelTemplates"),
          ("Output", "Downloadable, filled-in Excel school form file (.xlsx)")]
        ),
        ("7.0", "System Administration",
         "Manages user accounts, school settings, grading configuration, audit log monitoring, and manual sync triggers.",
         [("Input", "Admin commands (user create/update, settings changes, sync trigger)"),
          ("Reads", "D10: AuditLog"),
          ("Writes", "D1: Users, D2: Teachers, D11: SystemSettings, D12: GradingConfig, D10: AuditLog"),
          ("Output", "User list, current settings, audit log entries, sync status")]
        ),
    ]

    for num, name, desc, flows in processes:
        add_h2(doc, f"Process {num} — {name}")
        add_body(doc, desc)
        add_table(doc,
            ["Flow Type", "Detail"],
            flows,
            col_widths=[1.2, 5.5]
        )

    doc.add_page_break()

    # ── Section 5: Level 2 ──────────────────────────────────────────────────
    add_h1(doc, "5. Level 2 — Detailed Sub-Process Breakdowns")

    # --- 1.0 Auth
    add_h2(doc, "5.1  Process 1.0 — Authentication (Detail)")
    add_table(doc,
        ["Sub-Process", "Description", "Reads", "Writes"],
        [
            ["1.1 Validate Credentials", "Checks submitted username and hashed password against stored User record.", "D1: Users", "—"],
            ["1.2 Generate JWT Token", "Creates a signed JWT bearing the user's ID and role; returned to the browser.", "—", "— (token in memory)"],
            ["1.3 Log Login Activity", "Records the login event (user, IP, timestamp, role) in the audit log.", "—", "D10: AuditLog"],
        ],
        col_widths=[2.0, 3.2, 1.5, 1.5]
    )

    # --- 2.0 Sync
    add_h2(doc, "5.2  Process 2.0 — Data Synchronization (Detail)")
    add_table(doc,
        ["Sub-Process", "Description", "Source", "Writes To"],
        [
            ["2.1 Sync Teaching Loads", "Fetches faculty assignments from ATLAS API. Upserts Teachers, Subjects, Sections, and ClassAssignments.", "ATLAS API", "D2, D4, D5, D6"],
            ["2.2 Sync Enrollment Data", "Fetches students, sections, and enrollments from EnrollPro. Upserts Students, Sections, and Enrollments.", "EnrollPro API", "D3, D4, D7"],
            ["2.3 Record Sync Status", "Writes sync outcome (records matched/created/errors) to AuditLog. Admin can view via dashboard.", "Internal", "D10: AuditLog"],
        ],
        col_widths=[2.0, 3.0, 1.2, 1.5]
    )
    add_body(doc, "Sync runs automatically every 30 minutes. Admin can also trigger it manually via the Admin Dashboard.")

    # --- 3.0 Grades
    add_h2(doc, "5.3  Process 3.0 — Grade Management (Detail)")
    add_table(doc,
        ["Sub-Process", "Description", "Reads", "Writes"],
        [
            ["3.1 Record Scores", "Teacher submits Written Work scores, Performance Task scores, and Quarterly Assessment score per student.", "D6: ClassAssignments", "D8: Grades"],
            ["3.2 Compute Quarterly Grade", "Applies DepEd formula: (WW% × WW_PS) + (PT% × PT_PS) + (QA% × QA_PS). Weights come from GradingConfig.", "D8: Grades\nD12: GradingConfig", "D8: Grades (updated)"],
            ["3.3 Generate ECR", "Injects grade data into the ECR Excel template. Teacher downloads the filled class record.", "D8: Grades\nD3: Students\nD14: ECRTemplates", "— (file download)"],
        ],
        col_widths=[2.0, 3.0, 1.5, 1.5]
    )
    add_h2(doc, "DepEd Grading Formula")
    add_body(doc, "Quarterly Grade = (Written Work Weight × WW_PS) + (Performance Task Weight × PT_PS) + (Quarterly Assessment Weight × QA_PS)")
    add_table(doc,
        ["Subject Type", "Written Work", "Performance Task", "Quarterly Assessment"],
        [
            ["CORE (English, Filipino, AP, etc.)", "30%", "50%", "20%"],
            ["MATH / SCIENCE", "30%", "50%", "20%"],
            ["MAPEH", "30%", "50%", "20%"],
            ["TLE", "30%", "50%", "20%"],
        ],
        col_widths=[2.5, 1.3, 1.8, 1.8]
    )
    add_body(doc, "Note: Weights are configurable per subject type by the Admin via the Grading Config settings.")

    # --- 4.0 Attendance
    add_h2(doc, "5.4  Process 4.0 — Attendance Management (Detail)")
    add_table(doc,
        ["Sub-Process", "Description", "Reads", "Writes"],
        [
            ["4.1 Record Attendance", "Teacher marks each student's attendance status for a specific date: Present, Absent, Late, or Excused.", "D3: Students\nD4: Sections", "D9: Attendance"],
            ["4.2 Generate Attendance Reports", "Aggregates attendance records for a date range. Supplies data for SF2 and SF4 school form generation.", "D9: Attendance\nD3: Students\nD4: Sections", "— (report output)"],
        ],
        col_widths=[2.0, 3.2, 1.5, 1.5]
    )

    # --- 5.0 Templates
    add_h2(doc, "5.5  Process 5.0 — Template Management (Detail)")
    add_table(doc,
        ["Sub-Process", "Description", "Actor", "Writes To"],
        [
            ["5.1 Upload SF Template", "Registrar uploads an Excel (.xlsx) file for a specific school form (SF1–SF10). SMART parses and stores detected placeholder tags.", "Registrar", "D13: ExcelTemplates"],
            ["5.2 Upload ECR Template", "Teacher or Admin uploads an Excel (.xlsx) ECR template. Stored and linked to a subject type.", "Teacher / Admin", "D14: ECRTemplates"],
        ],
        col_widths=[2.0, 3.2, 1.3, 1.5]
    )

    # --- 6.0 Reports
    add_h2(doc, "5.6  Process 6.0 — Report Generation (Detail)")
    add_table(doc,
        ["Sub-Process", "Description", "Reads", "Output"],
        [
            ["6.1 Load Template", "Looks up the active Excel template for the requested form type from D13.", "D13: ExcelTemplates", "Template file in memory"],
            ["6.2 Populate with Real Data", "Replaces all {{PLACEHOLDER}} tags in the template with actual data (students, grades, attendance, school settings).", "D3, D4, D8, D9, D11", "Filled worksheet object"],
            ["6.3 Export File", "Converts the filled worksheet into a downloadable .xlsx file sent to the user's browser.", "—", "Downloaded .xlsx file"],
        ],
        col_widths=[2.0, 3.0, 1.5, 1.5]
    )

    add_h2(doc, "School Forms Supported")
    add_table(doc,
        ["Form", "Full Name", "Primary Data Used"],
        [
            ["SF1",  "School Register",                               "Students, Enrollments"],
            ["SF2",  "Daily Attendance Register",                     "Attendance"],
            ["SF3",  "Book of Accounts",                              "Grades"],
            ["SF4",  "Monthly Attendance Report",                     "Attendance"],
            ["SF5",  "Report on Promotion",                           "Grades"],
            ["SF6",  "Summarized Report on Promotion & Proficiency",  "Grades"],
            ["SF7",  "Homeroom Guidance Monitoring Sheet",            "Advisory / Sections"],
            ["SF8",  "Summary of Learner's Time on Task",             "Attendance"],
            ["SF9",  "Report Card",                                   "Grades"],
            ["SF10", "Permanent Record (Form 137)",                   "All — Students, Grades, Attendance"],
        ],
        col_widths=[0.8, 3.4, 2.5]
    )

    # --- 7.0 Admin
    add_h2(doc, "5.7  Process 7.0 — System Administration (Detail)")
    add_table(doc,
        ["Sub-Process", "Description", "Reads", "Writes"],
        [
            ["7.1 User Management", "Admin creates, updates, or disables user accounts. Teacher accounts are linked to Teacher profiles.", "D1: Users\nD2: Teachers", "D1: Users\nD2: Teachers"],
            ["7.2 School Settings", "Admin configures the active school year, current quarter, quarter date ranges, and grading component weights.", "D11, D12", "D11: SystemSettings\nD12: GradingConfig"],
            ["7.3 Audit Monitoring", "Admin views a filterable log of all significant system actions with timestamps, users, and severity levels.", "D10: AuditLog", "—"],
            ["7.4 Integration Control", "Admin manually triggers ATLAS or EnrollPro sync, or views the status and results of the last sync run.", "Sync state", "D10: AuditLog"],
        ],
        col_widths=[2.0, 3.0, 1.5, 1.5]
    )

    doc.add_page_break()

    # ── Section 6: Data Stores ───────────────────────────────────────────────
    add_h1(doc, "6. Data Stores Reference (Database Tables)")
    add_body(doc,
        "All data stores belong to SMART's local PostgreSQL database (smart_db). "
        "The schema is managed through Prisma ORM.")
    add_table(doc,
        ["ID", "Data Store", "Table Name", "Description"],
        [
            ["D1",  "Users",             "User",            "Login credentials (hashed password), role, name, email"],
            ["D2",  "Teachers",          "Teacher",         "Teacher profile, employee ID, linked to User account"],
            ["D3",  "Students",          "Student",         "LRN, full name, gender, birth date, guardian name & contact"],
            ["D4",  "Sections",          "Section",         "Section name, grade level (7–10), school year, assigned adviser"],
            ["D5",  "Subjects",          "Subject",         "Subject code, name, type (CORE/MAPEH/TLE/MATH_SCIENCE), grading weights"],
            ["D6",  "Class Assignments", "ClassAssignment", "Links Teacher → Subject → Section per school year; synced from ATLAS"],
            ["D7",  "Enrollments",       "Enrollment",      "Links Student → Section per school year; synced from EnrollPro"],
            ["D8",  "Grades",            "Grade",           "Per-student, per-quarter scores (WW, PT, QA) and computed final grade"],
            ["D9",  "Attendance",        "Attendance",      "Daily attendance per student per section (Present/Absent/Late/Excused)"],
            ["D10", "Audit Log",         "AuditLog",        "All system actions: user, action type, target, severity, timestamp"],
            ["D11", "System Settings",   "SystemSettings",  "School name, current school year/quarter, date ranges, UI colours"],
            ["D12", "Grading Config",    "GradingConfig",   "DepEd grading component weights per subject type"],
            ["D13", "Excel Templates",   "ExcelTemplate",   "Uploaded SF1–SF10 template files with placeholder maps"],
            ["D14", "ECR Templates",     "ECRTemplate",     "Uploaded Electronic Class Record template files per subject type"],
        ],
        col_widths=[0.5, 1.5, 1.8, 3.0]
    )

    doc.add_page_break()

    # ── Section 7: Data Flows ────────────────────────────────────────────────
    add_h1(doc, "7. Data Flows Reference")
    add_table(doc,
        ["#", "From", "To", "Data Description"],
        [
            ["F1",  "ATLAS",                   "Process 2.1",  "Faculty assignments: teacher ID, subject, section, school year"],
            ["F2",  "EnrollPro",               "Process 2.2",  "Students (LRN, name, gender), sections, enrollments, adviser mapping"],
            ["F3",  "Admin/Teacher/Registrar", "Process 1.1",  "Username, password, intended role"],
            ["F4",  "Process 1.2",             "User browser", "JWT token for session authorization"],
            ["F5",  "Teacher",                 "Process 3.1",  "Written Work scores, Performance Task scores, Quarterly Assessment score"],
            ["F6",  "Process 3.2",             "Teacher",      "Computed quarterly grade using DepEd formula"],
            ["F7",  "Teacher",                 "Process 4.1",  "Student ID, date, status (Present/Absent/Late/Excused)"],
            ["F8",  "Registrar",               "Process 5.1",  "Excel (.xlsx) SF template with {{PLACEHOLDER}} tags"],
            ["F9",  "Teacher/Admin",           "Process 5.2",  "Excel (.xlsx) ECR template"],
            ["F10", "Teacher/Registrar",       "Process 6.1",  "Form type, section ID, school year parameters"],
            ["F11", "Process 6.3",             "Teacher/Registrar", "Downloadable .xlsx file with real data populated"],
            ["F12", "Admin",                   "Process 7.1",  "New user data or role change command"],
            ["F13", "Admin",                   "Process 7.2",  "Updated school settings, grading weights, quarter dates"],
            ["F14", "Process 7.3",             "Admin",        "Filtered audit log entries (action, user, severity, time)"],
            ["F15", "Admin",                   "Process 7.4",  "Manual sync trigger command"],
            ["F16", "All Processes",           "D10: AuditLog","Action type, user ID, target, severity, timestamp"],
        ],
        col_widths=[0.4, 1.7, 1.7, 3.0]
    )

    doc.add_page_break()

    # ── Section 8: Tech Stack ────────────────────────────────────────────────
    add_h1(doc, "8. Architecture and Technology Stack")
    add_table(doc,
        ["Layer", "Technology"],
        [
            ["Frontend",        "React 18 + TypeScript + Vite + Tailwind CSS"],
            ["Backend",         "Node.js + Express + TypeScript"],
            ["Database",        "PostgreSQL (managed via Prisma ORM)"],
            ["Authentication",  "JWT Bearer Tokens (role-based: Admin / Teacher / Registrar)"],
            ["Excel Processing","xlsx-populate (template injection for school form generation)"],
            ["Network",         "Tailscale VPN (secure access to ATLAS and EnrollPro)"],
            ["Server Port",     "5004 (default)"],
        ],
        col_widths=[2.0, 4.5]
    )

    add_h2(doc, "8.1 Data Flow Constraints")
    constraints = [
        "ATLAS is READ-ONLY — SMART never sends data back to ATLAS.",
        "EnrollPro is READ-ONLY — SMART never sends data back to EnrollPro.",
        "All database writes go to smart_db (local PostgreSQL) only.",
        "Sync runs automatically every 30 minutes and can be manually triggered by Admin.",
        "JWT tokens are required for all API endpoints (except /api/auth/login).",
        "Each user role can only access pages and data within their authorization scope.",
    ]
    for c in constraints:
        add_bullet(doc, c)

    doc.add_page_break()

    # ── Footer note ──────────────────────────────────────────────────────────
    add_h1(doc, "Document Notes")
    add_body(doc,
        "This DFD documentation was generated by scanning the live SMART codebase, including:")
    for src in [
        "server/prisma/schema.prisma — database schema (all models and enums)",
        "server/src/index.ts — route registration and server entry point",
        "server/src/lib/atlasSync.ts — ATLAS synchronization service",
        "server/src/lib/enrollproSync.ts — EnrollPro synchronization service",
        "server/src/routes/ — all API route handlers",
    ]:
        add_bullet(doc, src)
    add_body(doc, "\nThis document reflects the system as of May 2026 (Capstone submission).")

    # ── Save ─────────────────────────────────────────────────────────────────
    out_path = os.path.join(os.path.dirname(__file__), "SMART_DFD.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
