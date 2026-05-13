#!/usr/bin/env python3
"""
SMART System – ERD Word Document Generator
Generates a professional ERD .docx with:
  1. Visual ERD diagram (matplotlib)
  2. Entity field detail tables
  3. Relationships summary table
  4. Enumeration types table
"""

import sys, os, subprocess

def ensure(pkg, import_name=None):
    name = import_name or pkg.replace("-", "_")
    try:
        __import__(name)
    except ImportError:
        print(f"  Installing {pkg}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])

print("Checking dependencies...")
ensure("matplotlib")
ensure("python-docx", "docx")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "SMART_ERD.docx")
IMG_MAIN    = os.path.join(SCRIPT_DIR, "_erd_main.png")
IMG_CFG     = os.path.join(SCRIPT_DIR, "_erd_cfg.png")

# ── Colour palette ────────────────────────────────────────────────────────────
C_USER   = "#1e4a7a"   # User / Teacher / SystemSettings
C_STUD   = "#1a5c8a"   # Student
C_SEC    = "#1a6b3a"   # Section / Enrollment / ExcelTemplate / ECRTemplate
C_SUB    = "#6b2d8b"   # Subject / Attendance / GradingConfig
C_JUN    = "#7a6100"   # ClassAssignment / Grade
C_AUDIT  = "#8b1a1a"   # AuditLog

# ── Entity definitions (diagram labels – abbreviated for readability) ─────────
MAIN_ENTITIES = {
    "User": {
        "color": C_USER,
        "fields": [
            ("PK", "id",        "String"),
            ("",   "username",  "String (unique)"),
            ("",   "password",  "String"),
            ("",   "role",      "Role"),
            ("",   "firstName", "String?"),
            ("",   "lastName",  "String?"),
            ("",   "email",     "String?"),
        ],
    },
    "Teacher": {
        "color": C_USER,
        "fields": [
            ("PK", "id",             "String"),
            ("FK", "userId",         "→ User"),
            ("",   "employeeId",     "String (unique)"),
            ("",   "specialization", "String?"),
        ],
    },
    "Student": {
        "color": C_STUD,
        "fields": [
            ("PK", "id",              "String"),
            ("",   "lrn",             "String (unique)"),
            ("",   "firstName",       "String"),
            ("",   "middleName",      "String?"),
            ("",   "lastName",        "String"),
            ("",   "suffix",          "String?"),
            ("",   "birthDate",       "DateTime?"),
            ("",   "gender",          "String?"),
            ("",   "guardianName",    "String?"),
        ],
    },
    "Section": {
        "color": C_SEC,
        "fields": [
            ("PK", "id",         "String"),
            ("",   "name",       "String"),
            ("",   "gradeLevel", "GradeLevel"),
            ("",   "schoolYear", "String"),
            ("FK", "adviserId",  "→ Teacher"),
        ],
    },
    "Subject": {
        "color": C_SUB,
        "fields": [
            ("PK", "id",               "String"),
            ("",   "code",             "String (unique)"),
            ("",   "name",             "String"),
            ("",   "type",             "SubjectType"),
            ("",   "writtenWork%",     "Int (30)"),
            ("",   "perfTask%",        "Int (50)"),
            ("",   "quarterlyAssess%", "Int (20)"),
        ],
    },
    "ClassAssignment": {
        "color": C_JUN,
        "fields": [
            ("PK", "id",             "String"),
            ("FK", "teacherId",      "→ Teacher"),
            ("FK", "subjectId",      "→ Subject"),
            ("FK", "sectionId",      "→ Section"),
            ("",   "schoolYear",     "String"),
            ("",   "ecrFileName",    "String?"),
        ],
    },
    "Enrollment": {
        "color": C_SEC,
        "fields": [
            ("PK", "id",         "String"),
            ("FK", "studentId",  "→ Student"),
            ("FK", "sectionId",  "→ Section"),
            ("",   "schoolYear", "String"),
            ("",   "status",     "String"),
        ],
    },
    "Grade": {
        "color": C_JUN,
        "fields": [
            ("PK", "id",                  "String"),
            ("FK", "studentId",           "→ Student"),
            ("FK", "classAssignmentId",   "→ ClassAssignment"),
            ("",   "quarter",             "Quarter"),
            ("",   "writtenWorkScores",   "Json?"),
            ("",   "perfTaskScores",      "Json?"),
            ("",   "quarterlyAssessScore","Float?"),
            ("",   "initialGrade",        "Float?"),
            ("",   "quarterlyGrade",      "Float?"),
            ("",   "remarks",             "String?"),
        ],
    },
    "Attendance": {
        "color": C_SUB,
        "fields": [
            ("PK", "id",         "String"),
            ("FK", "studentId",  "→ Student"),
            ("FK", "sectionId",  "→ Section"),
            ("",   "date",       "DateTime (Date)"),
            ("",   "status",     "AttendanceStatus"),
            ("",   "remarks",    "String?"),
            ("",   "recordedBy", "String?"),
        ],
    },
}

CFG_ENTITIES = {
    "AuditLog": {
        "color": C_AUDIT,
        "fields": [
            ("PK", "id",         "String"),
            ("",   "action",     "AuditAction"),
            ("",   "userId",     "String?"),
            ("",   "userName",   "String"),
            ("",   "userRole",   "String"),
            ("",   "target",     "String"),
            ("",   "targetType", "String"),
            ("",   "severity",   "AuditSeverity"),
            ("",   "ipAddress",  "String?"),
            ("",   "createdAt",  "DateTime"),
        ],
    },
    "SystemSettings": {
        "color": C_USER,
        "fields": [
            ("PK", "id",                "String ('main')"),
            ("",   "schoolName",        "String"),
            ("",   "division",          "String"),
            ("",   "region",            "String"),
            ("",   "currentSchoolYear", "String"),
            ("",   "currentQuarter",    "Quarter"),
            ("",   "logoUrl",           "String?"),
            ("",   "primaryColor",      "String"),
            ("",   "lastEnrollProSync", "DateTime?"),
        ],
    },
    "GradingConfig": {
        "color": C_SUB,
        "fields": [
            ("PK", "id",                    "String"),
            ("",   "subjectType",           "SubjectType (unique)"),
            ("",   "writtenWorkWeight",     "Int"),
            ("",   "performanceTaskWeight", "Int"),
            ("",   "quarterlyAssessWeight", "Int"),
            ("",   "isDepEdDefault",        "Boolean"),
        ],
    },
    "ExcelTemplate": {
        "color": C_SEC,
        "fields": [
            ("PK", "id",             "String"),
            ("",   "formType",       "FormType (unique)"),
            ("",   "formName",       "String"),
            ("",   "filePath",       "String"),
            ("",   "fileSize",       "Int"),
            ("",   "isActive",       "Boolean"),
            ("",   "uploadedBy",     "String"),
            ("",   "uploadedByName", "String"),
        ],
    },
    "ECRTemplate": {
        "color": C_SEC,
        "fields": [
            ("PK", "id",             "String"),
            ("",   "subjectName",    "String"),
            ("",   "subjectType",    "SubjectType?"),
            ("",   "filePath",       "String"),
            ("",   "fileSize",       "Int"),
            ("",   "isActive",       "Boolean"),
            ("",   "uploadedBy",     "String"),
            ("",   "uploadedByName", "String"),
        ],
    },
}

# ── Layout constants ──────────────────────────────────────────────────────────
BOX_W    = 3.1
ROW_H    = 0.28
HDR_H    = 0.44

def bh(entity_dict, name):
    """Box height for a named entity."""
    return HDR_H + ROW_H * len(entity_dict[name]["fields"])

def box_edge(positions, entity_dict, name, side):
    lx, ty = positions[name]
    h = bh(entity_dict, name)
    cx, cy = lx + BOX_W / 2, ty + h / 2
    if side == "left":   return (lx,          cy)
    if side == "right":  return (lx + BOX_W,  cy)
    if side == "top":    return (cx,           ty)
    if side == "bottom": return (cx,           ty + h)
    return (cx, cy)

# ── Diagram drawing ───────────────────────────────────────────────────────────
def draw_diagram(title, entities, positions, relationships,
                 canvas_w, canvas_h, img_path, dpi=150):
    """Render one ERD diagram and save to img_path."""

    fig, ax = plt.subplots(figsize=(canvas_w, canvas_h), dpi=dpi)
    ax.set_xlim(0, canvas_w)
    ax.set_ylim(0, canvas_h)
    ax.invert_yaxis()
    ax.axis("off")

    fig.patch.set_facecolor("#f5f6fa")
    ax.set_facecolor("#f5f6fa")

    # Title bar
    ax.add_patch(FancyBboxPatch((0.1, 0.08), canvas_w - 0.2, 0.50,
                                boxstyle="round,pad=0.03",
                                facecolor="#1e3a5f", edgecolor="none", zorder=1))
    ax.text(canvas_w / 2, 0.08 + 0.25, title,
            ha="center", va="center", fontsize=12, fontweight="bold",
            color="white", zorder=2)

    # Draw relationship lines first (behind boxes)
    for src, dst, src_side, dst_side, card, rad, label in relationships:
        p1 = box_edge(positions, entities, src, src_side)
        p2 = box_edge(positions, entities, dst, dst_side)
        ax.annotate("",
                    xy=p2, xytext=p1,
                    arrowprops=dict(
                        arrowstyle="-|>",
                        color="#7f8c8d",
                        lw=1.2,
                        mutation_scale=12,
                        connectionstyle=f"arc3,rad={rad}",
                    ), zorder=3)
        mx = (p1[0] + p2[0]) / 2
        my = (p1[1] + p2[1]) / 2
        ax.text(mx, my, card,
                ha="center", va="center", fontsize=6,
                color="#555",
                bbox=dict(boxstyle="round,pad=0.08",
                          facecolor="#f5f6fa", edgecolor="none", alpha=0.9),
                zorder=5)

    # Draw entity boxes
    for name, entity in entities.items():
        lx, ty = positions[name]
        h = bh(entities, name)
        color = entity["color"]

        # Shadow
        ax.add_patch(FancyBboxPatch(
            (lx + 0.05, ty + 0.07), BOX_W, h,
            boxstyle="round,pad=0.02",
            facecolor="#00000020", edgecolor="none", zorder=2))

        # White body
        ax.add_patch(FancyBboxPatch(
            (lx, ty), BOX_W, h,
            boxstyle="round,pad=0.02",
            facecolor="white", edgecolor=color, linewidth=1.4, zorder=3))

        # Coloured header
        ax.add_patch(FancyBboxPatch(
            (lx, ty), BOX_W, HDR_H,
            boxstyle="round,pad=0.02",
            facecolor=color, edgecolor=color, linewidth=1.4, zorder=4))

        ax.text(lx + BOX_W / 2, ty + HDR_H / 2, name,
                ha="center", va="center",
                fontsize=8, fontweight="bold", color="white", zorder=5)

        for i, (key, fname, ftype) in enumerate(entity["fields"]):
            fy = ty + HDR_H + ROW_H * i + ROW_H / 2

            # Zebra stripe
            if i % 2 == 0:
                ax.add_patch(FancyBboxPatch(
                    (lx + 0.01, ty + HDR_H + ROW_H * i), BOX_W - 0.02, ROW_H,
                    boxstyle="square,pad=0",
                    facecolor="#eef2ff", edgecolor="none", zorder=3))

            # PK / FK badge
            if key == "PK":
                ax.text(lx + 0.13, fy, "PK",
                        ha="center", va="center", fontsize=5.5,
                        color="white", fontweight="bold", zorder=6,
                        bbox=dict(boxstyle="round,pad=0.1",
                                  facecolor="#e67e22", edgecolor="none"))
            elif key == "FK":
                ax.text(lx + 0.13, fy, "FK",
                        ha="center", va="center", fontsize=5.5,
                        color="white", fontweight="bold", zorder=6,
                        bbox=dict(boxstyle="round,pad=0.1",
                                  facecolor="#2980b9", edgecolor="none"))

            ax.text(lx + 0.28, fy, fname,
                    ha="left", va="center", fontsize=6.5, zorder=6,
                    color="#111",
                    fontweight="bold" if key in ("PK", "FK") else "normal")
            ax.text(lx + BOX_W - 0.07, fy, ftype,
                    ha="right", va="center", fontsize=5.8, zorder=6,
                    color="#666", style="italic")

            # Row separator
            if i < len(entity["fields"]) - 1:
                ax.plot([lx + 0.05, lx + BOX_W - 0.05],
                        [ty + HDR_H + ROW_H * (i + 1)] * 2,
                        color="#ddd", lw=0.4, zorder=4)

    # Legend
    lx0, ly0 = canvas_w - 2.5, canvas_h - 2.2
    ax.add_patch(FancyBboxPatch(
        (lx0 - 0.1, ly0 - 0.1), 2.4, 1.8,
        boxstyle="round,pad=0.05",
        facecolor="white", edgecolor="#ccc", linewidth=0.8, zorder=6))
    ax.text(lx0 + 1.0, ly0 + 0.05, "LEGEND",
            ha="center", va="top", fontsize=7, fontweight="bold",
            color="#333", zorder=7)
    legend_items = [
        ("PK", "#e67e22", "Primary Key"),
        ("FK", "#2980b9", "Foreign Key"),
        ("1:N", "#7f8c8d", "One-to-Many"),
        ("1:1", "#7f8c8d", "One-to-One"),
    ]
    for j, (badge, c, desc) in enumerate(legend_items):
        by = ly0 + 0.38 + j * 0.32
        ax.text(lx0 + 0.18, by, badge,
                ha="center", va="center", fontsize=6, color="white",
                fontweight="bold", zorder=7,
                bbox=dict(boxstyle="round,pad=0.12", facecolor=c, edgecolor="none"))
        ax.text(lx0 + 0.42, by, desc,
                ha="left", va="center", fontsize=6.5, color="#333", zorder=7)

    plt.tight_layout(pad=0)
    plt.savefig(img_path, dpi=dpi, bbox_inches="tight",
                facecolor="#f5f6fa", format="png")
    plt.close()
    print(f"  Diagram saved → {img_path}")


# ── Diagram 1: Core academic entities ─────────────────────────────────────────
# Columns (x): 0.3 | 4.0 | 7.7 | 11.4
# Rows grow downward
MAIN_POS = {
    #  col1
    "User":             (0.3,  0.8),
    "Teacher":          (0.3,  3.55),
    "Student":          (0.3,  5.40),
    #  col2
    "Section":          (4.0,  0.8),
    "ClassAssignment":  (4.0,  3.00),
    "Enrollment":       (4.0,  5.40),
    "Attendance":       (4.0,  7.60),
    #  col3
    "Subject":          (7.7,  0.8),
    "Grade":            (7.7,  3.55),
}

# (src, dst, src_side, dst_side, card_label, arc_rad, text)
MAIN_REL = [
    ("User",            "Teacher",         "bottom", "top",    "1:1",  0.0,  ""),
    ("Teacher",         "Section",         "right",  "left",   "1:N",  0.1,  ""),
    ("Teacher",         "ClassAssignment", "right",  "left",   "1:N",  0.0,  ""),
    ("Subject",         "ClassAssignment", "left",   "right",  "1:N",  0.0,  ""),
    ("Section",         "ClassAssignment", "bottom", "top",    "1:N",  0.0,  ""),
    ("Section",         "Enrollment",      "right",  "top",    "1:N", -0.3,  ""),
    ("Section",         "Attendance",      "right",  "top",    "1:N", -0.5,  ""),
    ("Student",         "Enrollment",      "right",  "left",   "1:N",  0.0,  ""),
    ("Student",         "Grade",           "right",  "left",   "1:N",  0.3,  ""),
    ("Student",         "Attendance",      "right",  "left",   "1:N",  0.1,  ""),
    ("ClassAssignment", "Grade",           "right",  "left",   "1:N",  0.0,  ""),
]

# ── Diagram 2: Config / admin entities ────────────────────────────────────────
CFG_POS = {
    "AuditLog":       (0.3, 0.8),
    "SystemSettings": (4.0, 0.8),
    "GradingConfig":  (7.7, 0.8),
    "ExcelTemplate":  (0.3, 5.0),
    "ECRTemplate":    (4.0, 5.0),
}
CFG_REL = []   # cross-diagram links to User are documented in the summary table


# ── Word document helpers ─────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)

def add_hdr_row(table, labels, bg_hex):
    hdr = table.rows[0].cells
    for i, lbl in enumerate(labels):
        hdr[i].text = lbl
        r = hdr[i].paragraphs[0].runs[0]
        r.font.bold  = True
        r.font.size  = Pt(9)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        set_cell_bg(hdr[i], bg_hex)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


# ── Full field data for entity detail tables ──────────────────────────────────
FULL_FIELDS = {
    "User": [
        ("PK", "id",        "String",   "Unique identifier (CUID)"),
        ("",   "username",  "String",   "Unique login username"),
        ("",   "password",  "String",   "Bcrypt-hashed password"),
        ("",   "role",      "Role",     "TEACHER | ADMIN | REGISTRAR"),
        ("",   "firstName", "String?",  "First name"),
        ("",   "lastName",  "String?",  "Last name"),
        ("",   "email",     "String?",  "Email address"),
        ("",   "createdAt", "DateTime", "Record created timestamp"),
        ("",   "updatedAt", "DateTime", "Record last updated timestamp"),
    ],
    "Teacher": [
        ("PK", "id",             "String",   "Unique identifier"),
        ("FK", "userId",         "String",   "→ User.id  (cascade delete)"),
        ("",   "employeeId",     "String",   "DepEd Employee ID (unique)"),
        ("",   "specialization", "String?",  "Teaching specialization"),
        ("",   "createdAt",      "DateTime", "Record created timestamp"),
        ("",   "updatedAt",      "DateTime", "Record last updated timestamp"),
    ],
    "Student": [
        ("PK", "id",              "String",    "Unique identifier"),
        ("",   "lrn",             "String",    "Learner Reference Number (unique)"),
        ("",   "firstName",       "String",    "First name"),
        ("",   "middleName",      "String?",   "Middle name"),
        ("",   "lastName",        "String",    "Last name"),
        ("",   "suffix",          "String?",   "Name suffix (Jr., III, etc.)"),
        ("",   "birthDate",       "DateTime?", "Date of birth"),
        ("",   "gender",          "String?",   "Male / Female"),
        ("",   "address",         "String?",   "Home address"),
        ("",   "guardianName",    "String?",   "Parent / guardian full name"),
        ("",   "guardianContact", "String?",   "Parent / guardian contact number"),
        ("",   "createdAt",       "DateTime",  "Record created timestamp"),
        ("",   "updatedAt",       "DateTime",  "Record last updated timestamp"),
    ],
    "Section": [
        ("PK", "id",         "String",     "Unique identifier"),
        ("",   "name",       "String",     "Section name (e.g. 8-3 MAKAKALIKASAN)"),
        ("",   "gradeLevel", "GradeLevel", "GRADE_7 | GRADE_8 | GRADE_9 | GRADE_10"),
        ("",   "schoolYear", "String",     "School year label (e.g. 2026-2027)"),
        ("FK", "adviserId",  "String?",    "→ Teacher.id  (class adviser, nullable)"),
        ("",   "createdAt",  "DateTime",   "Record created timestamp"),
        ("",   "updatedAt",  "DateTime",   "Record last updated timestamp"),
        ("",   "[unique]",   "",           "UNIQUE (name, gradeLevel, schoolYear)"),
    ],
    "Subject": [
        ("PK", "id",                    "String",      "Unique identifier"),
        ("",   "code",                  "String",      "Subject code, unique (e.g. FIL7)"),
        ("",   "name",                  "String",      "Full subject name"),
        ("",   "description",           "String?",     "Optional description"),
        ("",   "type",                  "SubjectType", "CORE | MAPEH | TLE | MATH_SCIENCE"),
        ("",   "writtenWorkWeight",     "Int",         "Written work % weight (default 30)"),
        ("",   "perfTaskWeight",        "Int",         "Performance task % weight (default 50)"),
        ("",   "quarterlyAssessWeight", "Int",         "Quarterly assessment % weight (default 20)"),
        ("",   "createdAt",             "DateTime",    "Record created timestamp"),
        ("",   "updatedAt",             "DateTime",    "Record last updated timestamp"),
    ],
    "ClassAssignment": [
        ("PK", "id",              "String",    "Unique identifier"),
        ("FK", "teacherId",       "String",    "→ Teacher.id  (cascade delete)"),
        ("FK", "subjectId",       "String",    "→ Subject.id  (cascade delete)"),
        ("FK", "sectionId",       "String",    "→ Section.id  (cascade delete)"),
        ("",   "schoolYear",      "String",    "School year label (e.g. 2026-2027)"),
        ("",   "ecrFileName",     "String?",   "Linked ECR template filename"),
        ("",   "ecrLastSyncedAt", "DateTime?", "Last ECR template sync timestamp"),
        ("",   "createdAt",       "DateTime",  "Record created timestamp"),
        ("",   "updatedAt",       "DateTime",  "Record last updated timestamp"),
        ("",   "[unique]",        "",          "UNIQUE (teacherId, subjectId, sectionId, schoolYear)"),
    ],
    "Enrollment": [
        ("PK", "id",         "String",   "Unique identifier"),
        ("FK", "studentId",  "String",   "→ Student.id  (cascade delete)"),
        ("FK", "sectionId",  "String",   "→ Section.id  (cascade delete)"),
        ("",   "schoolYear", "String",   "School year label (e.g. 2026-2027)"),
        ("",   "status",     "String",   "ENROLLED | WITHDRAWN | TRANSFERRED | etc."),
        ("",   "createdAt",  "DateTime", "Record created timestamp"),
        ("",   "updatedAt",  "DateTime", "Record last updated timestamp"),
        ("",   "[unique]",   "",         "UNIQUE (studentId, sectionId, schoolYear)"),
    ],
    "Grade": [
        ("PK", "id",                    "String",   "Unique identifier"),
        ("FK", "studentId",             "String",   "→ Student.id  (cascade delete)"),
        ("FK", "classAssignmentId",     "String",   "→ ClassAssignment.id  (cascade delete)"),
        ("",   "quarter",               "Quarter",  "Q1 | Q2 | Q3 | Q4"),
        ("",   "writtenWorkScores",     "Json?",    "Array of {name, score, maxScore}"),
        ("",   "perfTaskScores",        "Json?",    "Array of {name, score, maxScore}"),
        ("",   "quarterlyAssessScore",  "Float?",   "Raw quarterly assessment score"),
        ("",   "quarterlyAssessMax",    "Float?",   "Max score for quarterly assessment (default 100)"),
        ("",   "writtenWorkPS",         "Float?",   "Written work percentage score"),
        ("",   "perfTaskPS",            "Float?",   "Performance task percentage score"),
        ("",   "quarterlyAssessPS",     "Float?",   "Quarterly assessment percentage score"),
        ("",   "initialGrade",          "Float?",   "Calculated initial grade"),
        ("",   "quarterlyGrade",        "Float?",   "Final quarterly grade (rounded)"),
        ("",   "remarks",               "String?",  "Passed | Failed | Incomplete | etc."),
        ("",   "createdAt",             "DateTime", "Record created timestamp"),
        ("",   "updatedAt",             "DateTime", "Record last updated timestamp"),
        ("",   "[unique]",              "",         "UNIQUE (studentId, classAssignmentId, quarter)"),
    ],
    "Attendance": [
        ("PK", "id",         "String",           "Unique identifier"),
        ("FK", "studentId",  "String",           "→ Student.id  (cascade delete)"),
        ("FK", "sectionId",  "String",           "→ Section.id  (cascade delete)"),
        ("",   "date",       "DateTime (Date)",  "Attendance date (date only, no time)"),
        ("",   "status",     "AttendanceStatus", "PRESENT | ABSENT | LATE | EXCUSED"),
        ("",   "remarks",    "String?",          "Optional remark / reason"),
        ("",   "recordedBy", "String?",          "User ID of teacher who recorded attendance"),
        ("",   "createdAt",  "DateTime",         "Record created timestamp"),
        ("",   "updatedAt",  "DateTime",         "Record last updated timestamp"),
        ("",   "[unique]",   "",                 "UNIQUE (studentId, sectionId, date)"),
    ],
    "AuditLog": [
        ("PK", "id",         "String",       "Unique identifier"),
        ("",   "action",     "AuditAction",  "CREATE | UPDATE | DELETE | LOGIN | LOGOUT | CONFIG"),
        ("FK", "userId",     "String?",      "Optional actor user ID (FK to User.id, null for system or legacy actions)"),
        ("",   "userName",   "String",       "Actor display name"),
        ("",   "userRole",   "String",       "Actor role at time of action"),
        ("",   "target",     "String",       "Human-readable description of affected entity"),
        ("",   "targetType", "String",       "Entity type (e.g. Student, Grade, Section)"),
        ("",   "targetId",   "String?",      "Affected record ID"),
        ("",   "details",    "String",       "Human-readable action description"),
        ("",   "ipAddress",  "String?",      "Client IP address"),
        ("",   "severity",   "AuditSeverity","INFO | WARNING | CRITICAL"),
        ("",   "metadata",   "Json?",        "Additional structured payload"),
        ("",   "createdAt",  "DateTime",     "Event timestamp"),
    ],
    "SystemSettings": [
        ("PK", "id",                 "String",    "Fixed primary key: 'main'"),
        ("",   "schoolName",         "String",    "Official school name"),
        ("",   "schoolId",           "String",    "DepEd school ID"),
        ("",   "division",           "String",    "DepEd division"),
        ("",   "region",             "String",    "DepEd region"),
        ("",   "address",            "String?",   "School address"),
        ("",   "contactNumber",      "String?",   "School contact number"),
        ("",   "email",              "String?",   "School email"),
        ("",   "currentSchoolYear",  "String",    "Active school year (e.g. 2026-2027)"),
        ("",   "currentQuarter",     "Quarter",   "Active grading quarter (Q1–Q4)"),
        ("",   "logoUrl",            "String?",   "Path to school logo file"),
        ("",   "primaryColor",       "String",    "Brand primary colour (hex)"),
        ("",   "secondaryColor",     "String",    "Brand secondary colour (hex)"),
        ("",   "accentColor",        "String",    "Brand accent colour (hex)"),
        ("",   "sessionTimeout",     "Int",       "Session timeout in minutes"),
        ("",   "maxLoginAttempts",   "Int",       "Max failed login attempts before lock"),
        ("",   "passwordMinLength",  "Int",       "Minimum password length"),
        ("",   "requireSpecialChar", "Boolean",   "Require special character in password"),
        ("",   "autoAdvanceQuarter", "Boolean",   "Auto-advance quarter on end date"),
        ("",   "lastEnrollProSync",  "DateTime?", "Last EnrollPro branding sync timestamp"),
        ("",   "q1StartDate",        "DateTime?", "Q1 start date"),
        ("",   "q1EndDate",          "DateTime?", "Q1 end date"),
        ("",   "q2StartDate",        "DateTime?", "Q2 start date"),
        ("",   "q2EndDate",          "DateTime?", "Q2 end date"),
        ("",   "q3StartDate",        "DateTime?", "Q3 start date"),
        ("",   "q3EndDate",          "DateTime?", "Q3 end date"),
        ("",   "q4StartDate",        "DateTime?", "Q4 start date"),
        ("",   "q4EndDate",          "DateTime?", "Q4 end date"),
        ("",   "createdAt",          "DateTime",  "Record created timestamp"),
        ("",   "updatedAt",          "DateTime",  "Record last updated timestamp"),
    ],
    "GradingConfig": [
        ("PK", "id",                    "String",      "Unique identifier"),
        ("",   "subjectType",           "SubjectType", "CORE | MAPEH | TLE | MATH_SCIENCE  (unique)"),
        ("",   "writtenWorkWeight",     "Int",         "Written work component % weight"),
        ("",   "performanceTaskWeight", "Int",         "Performance task component % weight"),
        ("",   "quarterlyAssessWeight", "Int",         "Quarterly assessment component % weight"),
        ("",   "isDepEdDefault",        "Boolean",     "True if using DepEd standard weights"),
        ("",   "createdAt",             "DateTime",    "Record created timestamp"),
        ("",   "updatedAt",             "DateTime",    "Record last updated timestamp"),
    ],
    "ExcelTemplate": [
        ("PK", "id",             "String",   "Unique identifier"),
        ("",   "formType",       "FormType", "SF1–SF10  (unique per form type)"),
        ("",   "formName",       "String",   "Display name (e.g. School Form 1)"),
        ("",   "description",    "String?",  "Template description"),
        ("",   "filePath",       "String",   "Absolute server-side file path"),
        ("",   "fileName",       "String",   "Original uploaded filename"),
        ("",   "fileSize",       "Int",      "File size in bytes"),
        ("",   "placeholders",   "Json?",    "Array of placeholder keys in the template"),
        ("",   "instructions",   "String?",  "Usage instructions for this template"),
        ("",   "isActive",       "Boolean",  "Whether template is active / in use"),
        ("",   "sheetName",      "String?",  "Target Excel sheet name for injection"),
        ("FK", "uploadedBy",     "String",   "Uploader user ID (FK to User.id)"),
        ("",   "uploadedByName", "String",   "Display name of uploader"),
        ("",   "createdAt",      "DateTime", "Record created timestamp"),
        ("",   "updatedAt",      "DateTime", "Record last updated timestamp"),
    ],
    "ECRTemplate": [
        ("PK", "id",             "String",      "Unique identifier"),
        ("",   "subjectName",    "String",      "Subject display name"),
        ("",   "description",    "String?",     "Template description"),
        ("",   "filePath",       "String",      "Absolute server-side file path"),
        ("",   "fileName",       "String",      "Original uploaded filename"),
        ("",   "fileSize",       "Int",         "File size in bytes"),
        ("",   "placeholders",   "Json?",       "Array of placeholder keys in the template"),
        ("",   "instructions",   "String?",     "Usage instructions"),
        ("",   "isActive",       "Boolean",     "Whether template is active / in use"),
        ("",   "subjectType",    "SubjectType?","Optional subject type classification"),
        ("FK", "uploadedBy",     "String",      "Uploader user ID (FK to User.id)"),
        ("",   "uploadedByName", "String",      "Display name of uploader"),
        ("",   "createdAt",      "DateTime",    "Record created timestamp"),
        ("",   "updatedAt",      "DateTime",    "Record last updated timestamp"),
    ],
}

# ── Relationships data for summary table ──────────────────────────────────────
ALL_RELATIONSHIPS = [
    ("User",            "1 : 1", "Teacher",         "User has one Teacher profile",          "Teacher.userId → User.id"),
    ("User",            "1 : N", "AuditLog",        "User can create many audit log entries", "AuditLog.userId → User.id"),
    ("User",            "1 : N", "ExcelTemplate",   "User can upload many SF templates",     "ExcelTemplate.uploadedBy → User.id"),
    ("User",            "1 : N", "ECRTemplate",     "User can upload many ECR templates",    "ECRTemplate.uploadedBy → User.id"),
    ("Teacher",         "1 : N", "Section",         "Teacher advises one or more Sections",  "Section.adviserId → Teacher.id"),
    ("Teacher",         "1 : N", "ClassAssignment", "Teacher has many Class Assignments",    "ClassAssignment.teacherId → Teacher.id"),
    ("Subject",         "1 : N", "ClassAssignment", "Subject used in many Class Assignments","ClassAssignment.subjectId → Subject.id"),
    ("Section",         "1 : N", "ClassAssignment", "Section has many Class Assignments",    "ClassAssignment.sectionId → Section.id"),
    ("Section",         "1 : N", "Enrollment",      "Section has many Enrollments",          "Enrollment.sectionId → Section.id"),
    ("Student",         "1 : N", "Enrollment",      "Student has many Enrollments",          "Enrollment.studentId → Student.id"),
    ("Student",         "1 : N", "Grade",           "Student has many Grade records",        "Grade.studentId → Student.id"),
    ("Student",         "1 : N", "Attendance",      "Student has many Attendance records",   "Attendance.studentId → Student.id"),
    ("Section",         "1 : N", "Attendance",      "Section has many Attendance records",   "Attendance.sectionId → Section.id"),
    ("ClassAssignment", "1 : N", "Grade",           "Class Assignment has many Grades",      "Grade.classAssignmentId → ClassAssignment.id"),
]

ENUMS = {
    "Role":             ["TEACHER", "ADMIN", "REGISTRAR"],
    "Quarter":          ["Q1", "Q2", "Q3", "Q4"],
    "GradeLevel":       ["GRADE_7", "GRADE_8", "GRADE_9", "GRADE_10"],
    "SubjectType":      ["CORE", "MAPEH", "TLE", "MATH_SCIENCE"],
    "AttendanceStatus": ["PRESENT", "ABSENT", "LATE", "EXCUSED"],
    "AuditAction":      ["CREATE", "UPDATE", "DELETE", "LOGIN", "LOGOUT", "CONFIG"],
    "AuditSeverity":    ["INFO", "WARNING", "CRITICAL"],
    "FormType":         ["SF1", "SF2", "SF3", "SF4", "SF5", "SF6", "SF7", "SF8", "SF9", "SF10"],
}

ENTITY_DESCRIPTIONS = {
    "User":             "System user accounts. All authenticated users (teacher, admin, registrar) have a User record.",
    "Teacher":          "Teacher profile linked 1-to-1 with a User. Holds the DepEd employee ID.",
    "Student":          "Student records identified by LRN (Learner Reference Number).",
    "Section":          "Class sections (e.g. Grade 8-3 MAKAKALIKASAN) scoped by school year.",
    "Subject":          "Academic subjects with DepEd grading component weights.",
    "ClassAssignment":  "Maps one teacher to one subject in one section for a given school year.",
    "Enrollment":       "Records a student's enrollment in a section for a given school year.",
    "Grade":            "Quarterly grade record per student per class assignment.",
    "Attendance":       "Daily attendance record per student per section.",
    "AuditLog":         "Immutable audit trail for all security-relevant actions in the system. Links to User when the actor still exists.",
    "SystemSettings":   "Singleton record (id='main') holding school-wide configuration.",
    "GradingConfig":    "DepEd grading component weights per subject type (overridable per school).",
    "ExcelTemplate":    "DepEd School Form templates (SF1–SF10) uploaded by a User for Excel export generation.",
    "ECRTemplate":      "Excel Class Record (ECR) templates per subject, uploaded by a User for grade export.",
}


# ── Build the Word document ───────────────────────────────────────────────────
def build_doc():
    doc = Document()

    # Page: A4 landscape for the whole document
    for sec in doc.sections:
        sec.page_width   = Cm(29.7)
        sec.page_height  = Cm(21.0)
        sec.top_margin   = Cm(1.5)
        sec.bottom_margin= Cm(1.5)
        sec.left_margin  = Cm(2.0)
        sec.right_margin = Cm(2.0)

    # ── Cover heading ───────────────────────────────────────────────────────
    t = doc.add_heading("SMART System", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(22)
    t.runs[0].font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    sub = doc.add_heading("Entity Relationship Diagram (ERD)", 2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    p = doc.add_paragraph("School Management and Reporting Tool  ·  smart_db  ·  PostgreSQL  ·  S.Y. 2026-2027")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # ── Section 1: Core ERD Diagram ─────────────────────────────────────────
    doc.add_heading("1.  Core Academic Entities — ERD Diagram", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(IMG_MAIN, width=Inches(10.8))

    doc.add_page_break()

    # ── Section 2: Config ERD Diagram ───────────────────────────────────────
    doc.add_heading("2.  Configuration & Administration Entities — ERD Diagram", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(IMG_CFG, width=Inches(8.5))

    doc.add_page_break()

    # ── Section 3: Entity Summary ────────────────────────────────────────────
    doc.add_heading("3.  Entity Summary", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    add_hdr_row(tbl, ["Entity", "Fields", "Description"], "1e3a5f")
    set_col_widths(tbl, [4.5, 1.5, 18.5])

    all_entities = {**MAIN_ENTITIES, **CFG_ENTITIES}
    for name in FULL_FIELDS:
        row = tbl.add_row().cells
        row[0].text = name
        row[0].paragraphs[0].runs[0].font.bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = str(len(FULL_FIELDS[name]))
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].paragraphs[0].runs[0].font.size = Pt(9)
        row[2].text = ENTITY_DESCRIPTIONS.get(name, "")
        row[2].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ── Section 4: Entity Field Details ─────────────────────────────────────
    doc.add_heading("4.  Entity Field Details", level=1)

    for name, fields in FULL_FIELDS.items():
        color = all_entities.get(name, {}).get("color", "#1e3a5f").lstrip("#")
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)

        h2 = doc.add_heading(name, level=2)
        h2.runs[0].font.color.rgb = RGBColor(r, g, b)

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        add_hdr_row(tbl, ["Key", "Field Name", "Type", "Description"], color)
        set_col_widths(tbl, [1.2, 4.5, 4.0, 14.8])

        for key, fname, ftype, desc in fields:
            row = tbl.add_row().cells

            row[0].text = key
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = row[0].paragraphs[0].runs[0]
            run0.font.size = Pt(8)
            run0.font.bold = True
            if key == "PK":
                run0.font.color.rgb = RGBColor(0xe6, 0x7e, 0x22)
            elif key == "FK":
                run0.font.color.rgb = RGBColor(0x29, 0x80, 0xb9)

            row[1].text = fname
            run1 = row[1].paragraphs[0].runs[0]
            run1.font.bold = key in ("PK", "FK")
            run1.font.size = Pt(8)

            row[2].text = ftype
            run2 = row[2].paragraphs[0].runs[0]
            run2.font.size = Pt(8)
            run2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
            run2.font.italic = True

            row[3].text = desc
            row[3].paragraphs[0].runs[0].font.size = Pt(8)

        doc.add_paragraph()

    doc.add_page_break()

    # ── Section 5: Relationships ─────────────────────────────────────────────
    doc.add_heading("5.  Relationships Summary", level=1)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    add_hdr_row(tbl, ["From Entity", "Cardinality", "To Entity", "Relationship Description", "FK Field"], "1e3a5f")
    set_col_widths(tbl, [3.5, 2.5, 3.5, 7.5, 7.5])

    for src, card, dst, desc, fk in ALL_RELATIONSHIPS:
        row = tbl.add_row().cells
        row[0].text = src
        row[0].paragraphs[0].runs[0].font.bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = card
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[1].paragraphs[0].runs[0].font.size = Pt(9)
        row[2].text = dst
        row[2].paragraphs[0].runs[0].font.bold = True
        row[2].paragraphs[0].runs[0].font.size = Pt(9)
        row[3].text = desc
        row[3].paragraphs[0].runs[0].font.size = Pt(9)
        row[4].text = fk
        run = row[4].paragraphs[0].runs[0]
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x29, 0x80, 0xb9)

    doc.add_page_break()

    # ── Section 6: Enumeration Types ─────────────────────────────────────────
    doc.add_heading("6.  Enumeration Types", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    add_hdr_row(tbl, ["Enum Name", "Allowed Values"], "1e3a5f")
    set_col_widths(tbl, [5.0, 19.5])

    for name, vals in ENUMS.items():
        row = tbl.add_row().cells
        row[0].text = name
        row[0].paragraphs[0].runs[0].font.bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = "  |  ".join(vals)
        row[1].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    doc.save(OUTPUT_PATH)
    print(f"  Word document saved → {OUTPUT_PATH}")


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("\n[ 1/4 ] Drawing Core ERD diagram...")
    draw_diagram(
        title="SMART – Core Academic Entities  (User · Teacher · Student · Section · Subject · ClassAssignment · Enrollment · Grade · Attendance)",
        entities=MAIN_ENTITIES,
        positions=MAIN_POS,
        relationships=MAIN_REL,
        canvas_w=15.0,
        canvas_h=13.0,
        img_path=IMG_MAIN,
        dpi=150,
    )

    print("[ 2/4 ] Drawing Config ERD diagram...")
    draw_diagram(
        title="SMART – Configuration & Administration Entities  (AuditLog · SystemSettings · GradingConfig · ExcelTemplate · ECRTemplate)",
        entities=CFG_ENTITIES,
        positions=CFG_POS,
        relationships=CFG_REL,
        canvas_w=12.0,
        canvas_h=9.5,
        img_path=IMG_CFG,
        dpi=150,
    )

    print("[ 3/4 ] Building Word document...")
    build_doc()

    print("[ 4/4 ] Cleaning up temp images...")
    for f in [IMG_MAIN, IMG_CFG]:
        if os.path.exists(f):
            os.remove(f)

    print(f"\nDone!  →  {OUTPUT_PATH}\n")
